import re
import os
import math
import urllib.request
from collections import Counter
from io import BytesIO

import matplotlib.pyplot as plt
import matplotlib
matplotlib.use("Agg")

from PIL import Image, ImageDraw
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


IMAGE_URL = "https://upload.wikimedia.org/wikipedia/commons/thumb/5/5a/Pride_and_Prejudice_-_Edwin_Austin_Abbey.jpg/640px-Pride_and_Prejudice_-_Edwin_Austin_Abbey.jpg"
LOGO_PATH = "logo.png"
OUTPUT_DOCX = "report.docx"
CHART_PATH = "chart.png"
COVER_IMG_PATH = "cover_image.png"

TITLE = "Pride and Prejudice"
AUTHOR = "Jane Austen"
SOURCE = "https://www.gutenberg.org/files/1342/1342-0.txt"

CHAPTER_TEXT = """Chapter I

It is a truth universally acknowledged, that a single man in possession of a good fortune, must be in want of a wife.

However little known the feelings or views of such a man may be on his first entering a neighbourhood, this truth is so well fixed in the minds of the surrounding families, that he is considered as the rightful property of some one or other of their daughters.

"My dear Mr. Bennet," said his lady to him one day, "have you heard that Netherfield Park is let at last?"

Mr. Bennet replied that he had not.

"But it is," returned she; "for Mrs. Long has just been here, and she told me all about it."

Mr. Bennet made no answer.

"Do not you want to know who has taken it?" cried his wife impatiently.

"You want to tell me, and I have no objection to hearing it."

This was invitation enough.

"Why, my dear, you must know, Mrs. Long says that Netherfield is taken by a young man of large fortune from the north of England; that he came down on Monday in a chaise and four to see the place, and was so much delighted with it, that he agreed with Mr. Morris immediately; that he is to take possession before Michaelmas, and some of his servants are to be in the house by the end of next week."

"What is his name?"

"Bingley."

"Is he married or single?"

"Oh! Single, my dear, to be sure! A single man of large fortune; four or five thousand a year. What a fine thing for our girls!"

"How so? How can it affect them?"

"My dear Mr. Bennet," replied his wife, "how can you be so tiresome! You must know that I am thinking of his marrying one of them."

"Is that his design in settling here?"

"Design! Nonsense, how can you talk so! But it is very likely that he may fall in love with one of them, and therefore you must visit him as soon as he comes."

"I see no occasion for that. You and the girls may go, or you may send them by themselves, which perhaps will be still better, for as you are as handsome as any of them, Mr. Bingley may like you the best of the party."

"My dear, you flatter me. I certainly have had my share of beauty, but I do not pretend to be anything extraordinary now. When a woman has five grown-up daughters, she ought to give over thinking of her own beauty."

"In such cases, a woman has not often much beauty to think of."

"But, my dear, you must indeed go and see Mr. Bingley when he comes into the neighbourhood."

"It is more than I engage for, I assure you."

"But consider your daughters. Only think what an establishment it would be for one of them. Sir William and Lady Lucas are determined to go, merely on that account, for in general, you know, they visit no newcomers. Indeed you must go, for it will be impossible for us to visit him if you do not."

"You are over-scrupulous, surely. I dare say Mr. Bingley will be very glad to see you; and I will send a few lines by you to assure him of my hearty consent to his marrying whichever he chooses of our girls; though I must throw in a good word for my little Lizzy."

"I desire you will do no such thing. Lizzy is not a bit better than the others; and I am sure she is not half so handsome as Jane, nor half so good-humoured as Lydia. But you are always giving her the preference."

"They have none of them much to recommend them," replied he; "they are all silly and ignorant like other girls; but Lizzy has something more of quickness than her sisters."

"Mr. Bennet, how can you abuse your own children so? You take delight in vexing me. You have no compassion on my poor nerves."

"You mistake me, my dear. I have a high respect for your nerves. They are my old friends. I have heard you mention them with consideration these twenty years at least."

"Ah, you do not know what I suffer."

"But I hope you will get over it, and live to see many young men of four thousand a year come into the neighbourhood."

"It will be no use to us, if twenty such should come, since you will not visit them."

"Depend upon it, my dear, that when there are twenty, I will visit them all."

Mr. Bennet was so odd a mixture of quick parts, sarcastic humour, reserve, and caprice, that the experience of three-and-twenty years had been insufficient to make his wife understand his character. Her mind was less difficult to develope. She was a woman of mean understanding, little information, and uncertain temper. When she was discontented, she fancied herself nervous. The business of her life was to get her daughters married; its solace was visiting and news."""


def count_words_per_paragraph(chapter_text):
    paragraphs = [p.strip() for p in chapter_text.split("\n\n") if p.strip()]
    counts = []
    for p in paragraphs:
        words = len(p.split())
        rounded = math.floor(words / 10) * 10
        counts.append(rounded)
    return counts, paragraphs


def make_chart(word_counts, path):
    freq = Counter(word_counts)
    keys = sorted(freq.keys())
    vals = [freq[k] for k in keys]

    fig, ax = plt.subplots(figsize=(8, 4))
    ax.bar([str(k) for k in keys], vals, color="#4c8bf5", edgecolor="white")
    ax.set_xlabel("Words per paragraph (rounded to 10)")
    ax.set_ylabel("Number of paragraphs")
    ax.set_title("Paragraph length distribution - Chapter I")
    fig.tight_layout()
    fig.savefig(path, dpi=150)
    plt.close(fig)
    print("Chart created.")


def download_image(url, path):
    headers = {"User-Agent": "Mozilla/5.0"}
    req = urllib.request.Request(url, headers=headers)
    with urllib.request.urlopen(req, timeout=20) as resp:
        data = resp.read()
    img = Image.open(BytesIO(data)).convert("RGB")
    img = img.crop((img.width // 4, 0, 3 * img.width // 4, img.height))
    img = img.resize((600, 400))
    img.save(path)


def make_cover_placeholder(path):
    img = Image.new("RGB", (600, 400), "#c8b89a")
    draw = ImageDraw.Draw(img)
    draw.rectangle([20, 20, 580, 380], outline="#5a3e2b", width=4)
    draw.text((180, 180), "Pride and Prejudice", fill="#2c1810")
    draw.text((220, 220), "Jane Austen", fill="#2c1810")
    img.save(path)


def make_logo(path):
    img = Image.new("RGB", (200, 200), "white")
    draw = ImageDraw.Draw(img)
    draw.ellipse([20, 20, 180, 180], outline="black", width=6)
    draw.rectangle([60, 60, 140, 140], outline="black", width=3)
    draw.text((70, 90), "P & P", fill="black")
    img = img.rotate(20)
    img.save(path)


def paste_logo_on_cover(cover_path, logo_path):
    cover = Image.open(cover_path).convert("RGBA")
    logo = Image.open(logo_path).convert("RGBA").resize((90, 90))
    cover.paste(logo, (cover.width - 110, cover.height - 110), logo)
    cover.convert("RGB").save(cover_path)


def build_report(word_counts, paragraphs, chart_path, cover_path):
    raw_counts = [len(p.split()) for p in paragraphs]
    total_words = sum(raw_counts)

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(11)

    h = doc.add_heading(TITLE, level=0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    h.runs[0].font.bold = True
    h.runs[0].font.size = Pt(28)

    if os.path.exists(cover_path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(cover_path, width=Inches(4))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"By {AUTHOR}")
    r.font.italic = True
    r.font.size = Pt(14)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Report by: Étudiant")
    r.font.size = Pt(12)

    doc.add_page_break()

    h2 = doc.add_heading("Paragraph length distribution", level=1)
    h2.runs[0].font.bold = True
    h2.runs[0].font.size = Pt(16)

    if os.path.exists(chart_path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(chart_path, width=Inches(5.5))

    doc.add_heading("Description", level=2)
    desc = (
        f"This report analyses the first chapter of '{TITLE}' by {AUTHOR}, "
        f"sourced from Project Gutenberg ({SOURCE}).\n\n"
        f"The chapter contains {len(paragraphs)} paragraphs and {total_words} words in total. "
        f"The shortest paragraph has {min(raw_counts)} words and the longest has {max(raw_counts)} words. "
        f"The average paragraph length is {sum(raw_counts)/len(raw_counts):.1f} words.\n\n"
        f"Word counts were rounded to the nearest ten for the distribution chart above."
    )
    doc.add_paragraph(desc)
    doc.save(OUTPUT_DOCX)
    print(f"Report saved: {OUTPUT_DOCX}")


def main():
    print("Processing book text...")
    word_counts, paragraphs = count_words_per_paragraph(CHAPTER_TEXT)

    make_chart(word_counts, CHART_PATH)

    print("Downloading cover image...")
    try:
        download_image(IMAGE_URL, COVER_IMG_PATH)
        print("Cover image downloaded.")
    except Exception as e:
        print(f"Image download failed ({e}), using placeholder.")
        make_cover_placeholder(COVER_IMG_PATH)

    make_logo(LOGO_PATH)
    paste_logo_on_cover(COVER_IMG_PATH, LOGO_PATH)
    print("Cover image ready.")

    build_report(word_counts, paragraphs, CHART_PATH, COVER_IMG_PATH)


if __name__ == "__main__":
    main()